# -*- coding: utf-8 -*-
"""
grade_analysis.py — دوال تحليل النتائج الدراسية (مشتركة بين الويب والمكتبة)
"""
import os, re, io, json, base64
from typing import List, Dict, Any, Optional

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

_GA_GRADES = [
    ("ممتاز",    90, 100,   "#27AE60", "#E8F8F0"),
    ("جيد جداً", 75, 89.99, "#2980B9", "#E8F4FD"),
    ("جيد",      65, 74.99, "#8E44AD", "#F3EAF9"),
    ("مقبول",    50, 64.99, "#E67E22", "#FEF3E2"),
    ("ضعيف",      0, 49.99, "#E74C3C", "#FDEDEC"),
]
_GA_SKIP_KW = frozenset([
    'اسم الطالب','رقم الهوية','الصف','القسم','الفصل','النظام الدراسي',
    'إشعار فترة أولى','إشعار','العام الدراسي','الفصل الدراسي',':','المادة',
    'مجموع','المجموع','النهاية العظمى','الختم','درجة المواظبة',
    'قائد المدرسة','السنة المشتركة','وزارة التعليم','منتظم',
])
_GA_SKIP_C = ['تقويم','مهام','اختبار','ملاحظة','درجة الس','المملكة','الإدارة']

# خريطة CID ← أسماء المواد (تُبنى تلقائياً من أول صفحة في PDF)
_GA_CID_MAP = {}

# الصفوف المعروفة في قالب تقارير نور (y coordinates)
_GA_NOOR_ROW_TOPS = [256, 273, 290, 307, 324, 341, 358, 375, 392]

# بيانات الترويسة القابلة للتعديل
_GA_HEADER_DATA = {
    "school":      "المدرسة",
    "school_year": "1446-1447",
    "term":        "الفصل الأول",
    "grade_class": "",
    "period":      "إشعار فترة أولى",
    "principal":   "",
    "teacher":     "",
    "region":      "الإدارة العامة للتعليم بالمنطقة",
}


def _ga_is_subject(v):
    if not v or v in _GA_SKIP_KW: return False
    if any(k in v for k in _GA_SKIP_C): return False
    return any('\u0600' <= ch <= '\u06FF' for ch in v)


def _ga_grade(pct):
    for lbl, lo, hi, col, bg in _GA_GRADES:
        if lo <= pct <= hi:
            return lbl, col, bg
    return "ضعيف", "#E74C3C", "#FDEDEC"


# ════════════════════════════════════════════════════════════
# ── محلّل PDF ذكي (نظام CID mapping) ─────────────────────
# ════════════════════════════════════════════════════════════
def _ga_build_cid_map(pdf):
    """
    يبني خريطة CID ← اسم مادة من الصفحة الأولى في PDF نور.
    يستخدم ملف Excel المصاحب (إذا وُجد) أو يعيد استخدام خريطة مبنية سابقاً.
    """
    global _GA_CID_MAP
    # استخدم الخريطة المبنية مسبقاً إذا كانت موجودة
    if _GA_CID_MAP:
        return _GA_CID_MAP

    # ابحث عن ملف Excel مصاحب لبناء أسماء المواد
    # إذا لا يوجد Excel، استخدم الأسماء الافتراضية من قالب نور
    default_subjects = [
        '*القرآن الكريم والتفسير',
        '*التربية الصحية والبدنية',
        'الأحياء',
        'الرياضيات',
        'اللغة الانجليزية',
        'الكفايات اللغوية',
        '*التفكير الناقد',
        'الكيمياء',
        '*التقنية الرقمية',
    ]

    chars0 = pdf.pages[0].chars
    cid_map = {}
    for i, (y_exp, subj) in enumerate(zip(_GA_NOOR_ROW_TOPS, default_subjects)):
        rc = [c for c in chars0 if abs(round(c['top']) - y_exp) <= 2 and c['x0'] > 240]
        if rc:
            key = tuple(c['text'] for c in sorted(rc, key=lambda x: x['x0']))
            cid_map[key] = subj

    _GA_CID_MAP = cid_map
    return cid_map


def _ga_parse_noor_pdf(filepath):
    """
    محلّل ذكي لملفات PDF نور — يستخدم CID mapping لاستخراج أسماء المواد.
    يعيد قائمة طلاب كاملة مع درجاتهم.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber غير مثبّت — pip install pdfplumber")

    def is_num(s):
        try:
            v = float(s)
            return 0 < v <= 1500
        except Exception:
            return False

    students = []

    with pdfplumber.open(filepath) as pdf:
        # بناء خريطة CID من الصفحة الأولى
        cid_map = _ga_build_cid_map(pdf)

        for pg_num, page in enumerate(pdf.pages):
            try:
                chars = page.chars
                words = page.extract_words()

                # استخراج رقم الهوية
                student_id = next((
                    w['text'] for w in words
                    if w['text'].isdigit() and 7 <= len(w['text']) <= 12
                    and 155 < w['top'] < 180
                ), None)

                # استخراج رقم الفصل (عادةً y≈112)
                section_num = next((
                    w['text'] for w in words
                    if w['text'].isdigit() and len(w['text']) == 1
                    and 108 < w['top'] < 120
                ), "1")

                # مطابقة المواد بخريطة CID
                subjects = []
                for y_exp in _GA_NOOR_ROW_TOPS:
                    rc = [c for c in chars
                          if abs(round(c['top']) - y_exp) <= 2 and c['x0'] > 240]
                    if not rc:
                        continue
                    key = tuple(c['text'] for c in sorted(rc, key=lambda x: x['x0']))
                    subj = cid_map.get(key)
                    if not subj:
                        continue

                    # استخراج الدرجات بالموضع X
                    row_nums = [
                        (w['x0'], float(w['text'])) for w in words
                        if not w['text'].startswith('(cid:')
                        and is_num(w['text'])
                        and abs(round(w['top']) - y_exp) <= 2
                    ]
                    mx = next((v for x, v in sorted(row_nums) if 25 < x < 35), None)
                    sc = next((v for x, v in sorted(row_nums) if 55 < x < 72), None)

                    if mx and sc is not None:
                        subjects.append({'subject': subj, 'score': sc, 'max_score': mx})

                if subjects:
                    students.append({
                        'name':        f'طالب_{pg_num + 1}',
                        'id':          student_id or str(pg_num + 1),
                        'class':       _GA_HEADER_DATA.get("grade_class", ""),
                        'section':     section_num,
                        'term':        _GA_HEADER_DATA.get("term", "الفصل الأول"),
                        'school_year': _GA_HEADER_DATA.get("school_year", "1446-1447"),
                        'school':      _GA_HEADER_DATA.get("school", ""),
                        'period':      _GA_HEADER_DATA.get("period", "إشعار فترة أولى"),
                        'subjects':    subjects,
                        'total_score': sum(s['score'] for s in subjects),
                        'total_max':   sum(s['max_score'] for s in subjects),
                    })
            except Exception as e:
                print(f"[GA-PDF] صفحة {pg_num}: {e}")
                continue

    return students


# ════════════════════════════════════════════════════════════
# ── محلّلات Excel و CSV ────────────────────────────────────
# ════════════════════════════════════════════════════════════
def _ga_parse_excel(filepath):
    students = []
    xl = pd.ExcelFile(filepath)
    for sheet_name in xl.sheet_names:
        try:
            df   = xl.parse(sheet_name, header=None)
            vals = df.values

            def find(kw):
                for r in range(len(vals)):
                    for c in range(len(vals[r])):
                        if kw in str(vals[r][c]).strip():
                            for cc in range(len(vals[r])):
                                v = str(vals[r][cc]).strip()
                                if v not in ['nan', '', kw, ':'] and v != str(vals[r][c]).strip():
                                    return v
                return None

            name = find('اسم الطالب')
            if not name or name == 'nan':
                continue

            school = _GA_HEADER_DATA.get("school") or None
            for r in range(min(12, len(vals))):
                for c in range(len(vals[r])):
                    v = str(vals[r][c]).strip()
                    if v not in ['nan', ''] and any(k in v for k in ['مدرسة', 'ثانوية', 'متوسطة', 'ابتدائية']):
                        school = v

            subjects = []
            total_score = total_max = 0
            for r in range(len(vals)):
                row = vals[r]
                sname = None
                for c in range(len(row) - 1, -1, -1):
                    v = str(row[c]).strip()
                    if v not in ['nan', ''] and not v.replace('.', '').replace('-', '').isdigit():
                        if _ga_is_subject(v):
                            sname = v
                        break
                if not sname:
                    continue
                try:
                    mx = float(str(row[1]).strip()) if str(row[1]).strip() not in ['nan', ''] else None
                    sc = float(str(row[2]).strip()) if str(row[2]).strip() not in ['nan', ''] else None
                except Exception:
                    mx = sc = None
                if mx and sc is not None and mx > 0:
                    if 'المجموع' in sname or sname == 'مجموع':
                        total_score, total_max = sc, mx
                    else:
                        subjects.append({'subject': sname, 'score': sc, 'max_score': mx})

            if subjects:
                gc = find('الصف') or _GA_HEADER_DATA.get("grade_class", "")
                tm = find('الفصل الدراسي') or _GA_HEADER_DATA.get("term", "")
                sy = find('العام الدراسي') or _GA_HEADER_DATA.get("school_year", "")
                pr = find('إشعار') or _GA_HEADER_DATA.get("period", "إشعار فترة أولى")
                students.append({
                    'name':        name,
                    'id':          find('رقم الهوية') or '',
                    'class':       gc,
                    'section':     find('الفصل') or '',
                    'term':        tm,
                    'school_year': sy,
                    'school':      school or '',
                    'period':      pr,
                    'subjects':    subjects,
                    'total_score': total_score or sum(s['score'] for s in subjects),
                    'total_max':   total_max   or sum(s['max_score'] for s in subjects),
                })
        except Exception as e:
            print(f"[GA-XLS] ورقة {sheet_name}: {e}")
    return students


def _ga_parse_csv(filepath):
    df   = pd.read_csv(filepath, encoding='utf-8-sig')
    cols = list(df.columns)
    name_col    = next((c for c in cols if 'اسم' in c or 'طالب' in c), cols[0] if cols else None)
    class_col   = next((c for c in cols if 'صف' in c or 'فصل' in c), None)
    subject_col = next((c for c in cols if 'مادة' in c), None)
    score_col   = next((c for c in cols if 'درجة' in c or 'نتيجة' in c), None)
    max_col     = next((c for c in cols if 'نهاية' in c or 'أعلى' in c), None)
    if not name_col:
        raise ValueError("لم يُعثر على عمود اسم الطالب")
    students = {}
    if subject_col and score_col:
        for _, row in df.iterrows():
            n = str(row[name_col]).strip()
            if not n or n == 'nan':
                continue
            if n not in students:
                students[n] = {
                    'name': n, 'id': '',
                    'class': str(row[class_col]).strip() if class_col else _GA_HEADER_DATA.get("grade_class", ""),
                    'section': '', 'term': _GA_HEADER_DATA.get("term", ""),
                    'school_year': _GA_HEADER_DATA.get("school_year", ""),
                    'school': _GA_HEADER_DATA.get("school", ""),
                    'period': _GA_HEADER_DATA.get("period", "إشعار"),
                    'subjects': [], 'total_score': 0, 'total_max': 0,
                }
            try:
                sc = float(row[score_col])
            except Exception:
                sc = 0
            try:
                mx = float(row[max_col]) if max_col else 100
            except Exception:
                mx = 100
            students[n]['subjects'].append({
                'subject': str(row[subject_col]).strip(), 'score': sc, 'max_score': mx
            })
    else:
        fixed = {name_col, class_col}
        scols = [c for c in cols if c not in fixed]
        for _, row in df.iterrows():
            n = str(row[name_col]).strip()
            if not n or n == 'nan':
                continue
            subjs = []
            for sc_col in scols:
                try:
                    subjs.append({'subject': sc_col, 'score': float(row[sc_col]), 'max_score': 100})
                except Exception:
                    pass
            if subjs:
                students[n] = {
                    'name': n, 'id': '',
                    'class': str(row[class_col]).strip() if class_col else _GA_HEADER_DATA.get("grade_class", ""),
                    'section': '', 'term': _GA_HEADER_DATA.get("term", ""),
                    'school_year': _GA_HEADER_DATA.get("school_year", ""),
                    'school': _GA_HEADER_DATA.get("school", ""), 'period': 'إشعار',
                    'subjects': subjs, 'total_score': 0, 'total_max': 0,
                }
    result = list(students.values())
    for s in result:
        s['total_score'] = sum(x['score'] for x in s['subjects'])
        s['total_max']   = sum(x['max_score'] for x in s['subjects'])
    return result


def _ga_parse_file(filepath):
    global _GA_CID_MAP
    ext = os.path.splitext(filepath)[1].lower()
    if ext in (".xlsx", ".xls"):
        _GA_CID_MAP = {}  # reset for new file
        return _ga_parse_excel(filepath)
    elif ext == ".csv":
        return _ga_parse_csv(filepath)
    elif ext == ".pdf":
        _GA_CID_MAP = {}  # reset to rebuild from this PDF
        return _ga_parse_noor_pdf(filepath)
    else:
        raise ValueError(f"صيغة غير مدعومة: {ext}\nالصيغ المدعومة: xlsx, xls, pdf, csv")


# ════════════════════════════════════════════════════════════
# ── HTML للعرض الداخلي ────────────────────────────────────
# ════════════════════════════════════════════════════════════
def _ga_build_html(students, sel_subject="الكل"):
    if not students:
        return _ga_placeholder_html("لا توجد بيانات للعرض")

    first = students[0]
    subjects_list = (
        sorted(set(sub["subject"] for s in students for sub in s.get("subjects", [])))
        if sel_subject == "الكل" else [sel_subject]
    )

    # إحصائيات الشريط الجانبي
    all_pcts = []
    for s in students:
        for sub in s.get("subjects", []):
            if sel_subject in ("الكل", sub["subject"]) and sub["max_score"] > 0:
                all_pcts.append(sub["score"] / sub["max_score"] * 100)
    avg_all   = sum(all_pcts) / len(all_pcts) if all_pcts else 0
    pass_rate = sum(1 for p in all_pcts if p >= 50) / len(all_pcts) * 100 if all_pcts else 0
    grade_all = {g[0]: 0 for g in _GA_GRADES}
    for p in all_pcts:
        lbl, *_ = _ga_grade(p)
        grade_all[lbl] += 1

    hdr = _GA_HEADER_DATA
    sidebar = f"""\n<div style="width:205px;min-width:205px;background:#F8FBFF;
                border-left:1px solid #DDE3EC;padding:14px 10px;
                font-family:Tahoma,Arial,sans-serif;font-size:12px;
                direction:rtl;overflow-y:auto;height:100%;box-sizing:border-box;">
      <div style="font-weight:bold;color:#1A3A5C;font-size:13px;margin-bottom:10px;
                  border-bottom:2px solid #1A3A5C;padding-bottom:5px;">📋 ملخص النتائج</div>
      <div style="background:#EBF5FB;border-radius:8px;padding:8px 10px;margin-bottom:7px;">
        <div style="color:#7F8C8D;font-size:10px;">عدد الطلاب</div>
        <div style="color:#1A3A5C;font-size:22px;font-weight:900;">{len(students)}</div>
      </div>
      <div style="background:#EBF5FB;border-radius:8px;padding:8px 10px;margin-bottom:7px;">
        <div style="color:#7F8C8D;font-size:10px;">متوسط التحصيل</div>
        <div style="color:#2471A3;font-size:20px;font-weight:900;">{avg_all:.1f}%</div>
      </div>
      <div style="background:{'#E8F8F0' if pass_rate>=70 else '#FEF3E2'};border-radius:8px;
                  padding:8px 10px;margin-bottom:10px;">
        <div style="color:#7F8C8D;font-size:10px;">نسبة النجاح</div>
        <div style="color:{'#27AE60' if pass_rate>=70 else '#E67E22'};font-size:20px;font-weight:900;">{pass_rate:.1f}%</div>
        <div style="background:#D5D8DC;border-radius:4px;height:6px;margin-top:4px;">
          <div style="background:{'#27AE60' if pass_rate>=70 else '#E67E22'};
                      width:{min(pass_rate,100):.0f}%;height:6px;border-radius:4px;"></div>
        </div>
      </div>
      <div style="font-weight:bold;color:#555;font-size:10px;margin-bottom:6px;">توزيع التقديرات</div>
      {"".join(f'''<div style="display:flex;justify-content:space-between;align-items:center;
            background:{bg};border-radius:6px;padding:5px 8px;margin-bottom:4px;">
        <span style="color:{col};font-weight:bold;font-size:15px;">{grade_all[lbl]}</span>
        <div><div style="color:{col};font-size:11px;font-weight:bold;">{lbl}</div>
        <div style="color:#999;font-size:9px;">{(grade_all[lbl]/len(all_pcts)*100 if all_pcts else 0):.0f}%</div></div>
      </div>''' for lbl,_,_,col,bg in _GA_GRADES)}
      <div style="margin-top:12px;padding-top:10px;border-top:1px solid #DDE3EC;">
        <div style="font-size:9px;color:#7F8C8D;font-weight:bold;margin-bottom:4px;">بيانات التقرير</div>
        <div style="font-size:10px;color:#1A3A5C;line-height:1.8;">
          <b>المدرسة:</b> {hdr.get('school','—')}<br>
          <b>الصف:</b> {first.get('class','') or hdr.get('grade_class','—')}<br>
          <b>الفصل:</b> {first.get('term','') or hdr.get('term','—')}<br>
          <b>العام:</b> {first.get('school_year','') or hdr.get('school_year','—')}هـ
        </div>
      </div>
    </div>"""

    # بناء كروت المواد
    cards = ""
    for subj_name in subjects_list:
        pairs = []
        for s in students:
            for sub in s.get("subjects", []):
                if sub["subject"] == subj_name:
                    pairs.append((sub["score"], sub["max_score"]))
                    break
        if not pairs:
            continue
        n = len(pairs)
        max_score = pairs[0][1]
        pcts = [sc / mx * 100 if mx > 0 else 0 for sc, mx in pairs]
        avg  = sum(pcts) / len(pcts)
        high = max(sc for sc, mx in pairs)
        low  = min(sc for sc, mx in pairs)
        gc = {g[0]: 0 for g in _GA_GRADES}
        for p in pcts:
            lbl, *_ = _ga_grade(p)
            gc[lbl] += 1

        rows_html = "".join(f"""\n<tr style="background:{'#F8FBFF' if i%2==0 else 'white'}">
              <td style="padding:7px 10px;text-align:center;">
                <span style="background:{col};color:white;border-radius:10px;
                             padding:2px 12px;font-size:11px;font-weight:bold;">{lbl}</span>
              </td>
              <td style="padding:7px;text-align:center;color:#7F8C8D;font-size:11px;">
                {int(lo)}–{int(hi) if hi<100 else int(max_score)}</td>
              <td style="padding:7px;text-align:center;font-weight:bold;color:{col};font-size:14px;">{gc[lbl]}</td>
              <td style="padding:7px;text-align:center;color:#7F8C8D;font-size:11px;">
                {gc[lbl]/n*100:.0f}%</td>
            </tr>""" for i, (lbl, lo, hi, col, bg) in enumerate(_GA_GRADES))

        bars = "".join(f"""\n<div style="display:flex;align-items:center;gap:6px;margin-bottom:5px;">
              <div style="width:52px;text-align:right;font-size:10px;color:{col};font-weight:bold;">{lbl}</div>
              <div style="flex:1;background:#F0F0F0;border-radius:4px;height:18px;">
                <div style="background:{col};width:{gc[lbl]/n*100:.0f}%;height:18px;
                             border-radius:4px;min-width:{'2px' if gc[lbl]>0 else '0'};"></div>
              </div>
              <div style="width:28px;font-size:11px;font-weight:bold;color:{col};">{gc[lbl]}</div>
            </div>""" for lbl, _, _, col, _ in _GA_GRADES)

        cards += f"""\n<div style="border:1.5px solid #DDE3EC;border-radius:12px;margin-bottom:18px;
                    overflow:hidden;box-shadow:0 2px 12px rgba(26,58,92,0.07);">
          <div style="background:linear-gradient(135deg,#0F2542,#1A3A5C);padding:13px 20px;
                      display:flex;align-items:center;justify-content:space-between;">
            <span style="color:white;font-size:14px;font-weight:bold;">📚 {subj_name}</span>
            <span style="background:rgba(255,255,255,0.18);color:white;border-radius:10px;
                         padding:2px 12px;font-size:11px;">{n} طالب</span>
          </div>
          <div style="display:flex;background:#F5F8FE;border-bottom:1px solid #DDE3EC;">
            {"".join(f'''<div style="flex:1;text-align:center;padding:10px 4px;border-left:1px solid #DDE3EC;">
              <div style="font-size:18px;font-weight:900;color:{c};">{v}</div>
              <div style="font-size:10px;color:#7F8C8D;">{l}</div>
            </div>''' for l,v,c in [
                ("المتوسط",    f"{avg:.1f}%",   "#2471A3"),
                ("أعلى درجة", f"{high:.0f}",    "#27AE60"),
                ("أقل درجة",  f"{low:.0f}",     "#E74C3C"),
                ("الناجحون",  str(sum(1 for p in pcts if p>=50)), "#8E44AD"),
            ])}
          </div>
          <div style="display:flex;">
            <div style="flex:1;padding:12px;border-left:1px solid #DDE3EC;">
              <table style="width:100%;border-collapse:collapse;font-family:Tahoma,Arial,sans-serif;">
                <thead><tr style="background:#1A3A5C;">
                  <th style="padding:7px;color:white;font-size:11px;">المستوى</th>
                  <th style="padding:7px;color:white;font-size:11px;">النطاق</th>
                  <th style="padding:7px;color:white;font-size:11px;">العدد</th>
                  <th style="padding:7px;color:white;font-size:11px;">%</th>
                </tr></thead>
                <tbody>{rows_html}</tbody>
              </table>
            </div>
            <div style="flex:1;padding:14px 16px;background:#FAFCFF;">
              <div style="font-size:10px;color:#7F8C8D;font-weight:bold;margin-bottom:10px;">التوزيع البصري</div>
              {bars}
            </div>
          </div>
        </div>"""

    meta_bar = f"{first.get('class','') or ''} — {first.get('term','') or ''} {first.get('school_year','') or ''}هـ — {first.get('school','') or ''}"
    return f"""<!DOCTYPE html><html lang="ar" dir="rtl"><head><meta charset="UTF-8">\n<style>*{{box-sizing:border-box;margin:0;padding:0;}}\nbody{{font-family:Tahoma,Arial,sans-serif;background:#F0F4F8;direction:rtl;color:#1A2332;}}\n::-webkit-scrollbar{{width:5px;}}::-webkit-scrollbar-thumb{{background:#BDC3C7;border-radius:3px;}}\n</style></head><body>\n<div style="display:flex;height:100vh;overflow:hidden;">
  {sidebar}
  <div style="flex:1;overflow-y:auto;padding:16px;">
    <div style="color:#1A3A5C;font-size:11px;font-weight:bold;margin-bottom:14px;
                padding-bottom:8px;border-bottom:2px solid #DDE3EC;">
      📊 نتائج التحليل — {meta_bar}
    </div>
    {cards if cards else '<div style="text-align:center;color:#7F8C8D;padding:60px;font-size:14px;">لا توجد نتائج</div>'}
  </div>
</div></body></html>"""


# ════════════════════════════════════════════════════════════
# ── HTML للطباعة/PDF ──────────────────────────────────────
# ════════════════════════════════════════════════════════════
def _ga_build_print_html(students, sel_subject="الكل"):
    if not students:
        return "<html><body><p>لا توجد بيانات</p></body></html>"
    first = students[0]
    hdr   = _GA_HEADER_DATA
    subjects_list = (
        sorted(set(sub["subject"] for s in students for sub in s.get("subjects", [])))
        if sel_subject == "الكل" else [sel_subject]
    )
    pages = ""
    for subj_name in subjects_list:
        pairs = []
        for s in students:
            for sub in s.get("subjects", []):
                if sub["subject"] == subj_name:
                    pairs.append((sub["score"], sub["max_score"]))
                    break
        if not pairs:
            continue
        n  = len(pairs)
        mx = pairs[0][1]
        pcts  = [sc / m * 100 if m > 0 else 0 for sc, m in pairs]
        avg   = sum(pcts) / len(pcts)
        high  = max(sc for sc, m in pairs)
        low   = min(sc for sc, m in pairs)
        total = sum(sc for sc, m in pairs)
        gc = {g[0]: 0 for g in _GA_GRADES}
        for p in pcts:
            lbl, *_ = _ga_grade(p)
            gc[lbl] += 1

        cl  = first.get('class', '') or hdr.get('grade_class', '—')
        tm  = first.get('term', '')  or hdr.get('term', '—')
        sy  = first.get('school_year', '') or hdr.get('school_year', '—')
        sch = first.get('school', '') or hdr.get('school', '—')
        pr  = first.get('period', '') or hdr.get('period', 'الاختبار')
        reg = hdr.get('region', 'الإدارة العامة للتعليم بالمنطقة')
        tchr = hdr.get('teacher', '')
        prin = hdr.get('principal', '')

        grade_rows = "".join(f"""\n<tr>\n<td style="padding:7px 10px;border-bottom:1px solid #E8ECF0;text-align:center;">
                <span style="background:{col};color:white;border-radius:10px;
                             padding:2px 14px;font-size:10px;font-weight:bold;">{lbl}</span></td>
              <td style="padding:7px;border-bottom:1px solid #E8ECF0;text-align:center;
                          color:#7F8C8D;font-size:11px;">{int(lo)}–{int(hi) if hi<100 else int(mx)}</td>
              <td style="padding:7px;border-bottom:1px solid #E8ECF0;text-align:center;
                          font-weight:bold;color:{col};font-size:13px;">{gc[lbl]}</td>
              <td style="padding:7px;border-bottom:1px solid #E8ECF0;text-align:center;
                          color:#7F8C8D;font-size:11px;">{gc[lbl]/n*100:.0f}%</td>
            </tr>""" for lbl, lo, hi, col, bg in _GA_GRADES)

        bar_max = max(gc.values()) or 1
        bars_vis = "".join(f"""\n<div style="display:flex;align-items:center;gap:6px;margin-bottom:6px;">
              <span style="min-width:52px;text-align:right;font-size:10px;
                            color:{col};font-weight:bold;">{lbl}</span>
              <div style="width:130px;background:#F0F0F0;border-radius:3px;height:16px;">
                <div style="background:{col};width:{int(gc[lbl]/bar_max*130)}px;
                             height:16px;border-radius:3px;"></div>
              </div>
              <span style="font-size:11px;font-weight:bold;color:{col};">{gc[lbl]}</span>
              <span style="font-size:9px;color:#999;">({gc[lbl]/n*100:.0f}%)</span>
            </div>""" for lbl, _, _, col, _ in _GA_GRADES)

        donuts = "".join(f"""\n<div style="text-align:center;width:70px;">
              <div style="width:60px;height:60px;border-radius:50%;
                  background:conic-gradient({col} {gc[lbl]/n*360 if n>0 else 0:.0f}deg,#E8ECF0 0deg);
                  margin:0 auto 4px;display:flex;align-items:center;justify-content:center;">
                <div style="width:40px;height:40px;background:white;border-radius:50%;
                             display:flex;align-items:center;justify-content:center;
                             font-size:10px;font-weight:bold;color:{col};">
                  {gc[lbl]/n*100:.0f}%</div>
              </div>
              <div style="font-size:9px;font-weight:bold;color:{col};">{lbl}</div>
            </div>""" for lbl, _, _, col, _ in _GA_GRADES if gc[lbl] > 0)

        pages += f"""\n<div class="page">
          <div style="background:linear-gradient(135deg,#0F2542,#1A3A5C);padding:18px 28px 14px;">
            <div style="display:flex;align-items:flex-start;justify-content:space-between;">
              <div style="text-align:right;">
                <div style="font-size:9px;color:rgba(255,255,255,0.7);">{reg}</div>
                <div style="font-size:11px;font-weight:bold;color:white;">{sch}</div>
                <div style="font-size:9px;color:rgba(255,255,255,0.6);margin-top:2px;">مدرسة</div>
              </div>
              <div style="text-align:center;">
                <div style="font-size:8px;letter-spacing:4px;color:rgba(255,255,255,0.5);margin-bottom:4px;">● ● ● ● ●</div>
                <div style="font-size:13px;font-weight:bold;color:white;">وزارة التعليم</div>
                <div style="font-size:8px;color:rgba(255,255,255,0.55);">Ministry of Education</div>
              </div>
              <div style="font-size:9px;color:rgba(255,255,255,0.8);">{pr}</div>
            </div>
            <div style="background:rgba(255,255,255,0.12);border:1px solid rgba(255,255,255,0.25);
                        border-radius:8px;margin-top:12px;padding:9px 0;text-align:center;">
              <span style="font-size:15px;font-weight:bold;color:white;">
                تحليل نتائج مادة &nbsp;[&nbsp;{subj_name}&nbsp;]
              </span>
            </div>
          </div>
          <div style="display:flex;background:#F5F8FE;border-bottom:2px solid #DDE3EC;">
            {"".join(f'''<div style="flex:1;text-align:center;padding:10px;border-left:1px solid #DDE3EC;">
              <div style="font-size:8px;color:#7F8C8D;font-weight:bold;">{l}</div>
              <div style="font-size:13px;font-weight:bold;color:#1A3A5C;margin-top:3px;">{v}</div>
            </div>''' for l,v in [("المرحلة / الصف",cl),("السنة / الفصل",f"{tm} / {sy}هـ"),("درجة القياس",str(int(mx)))])}
          </div>
          <div style="text-align:center;padding:12px 0 6px;">
            <span style="background:linear-gradient(135deg,#2471A3,#1A3A5C);color:white;
                         border-radius:16px;padding:4px 24px;font-size:11px;font-weight:bold;">
              الإحصائيات التفصيلية
            </span>
          </div>
          <div style="display:flex;gap:0;margin:0 20px;">
            <div style="width:190px;min-width:190px;background:#F8FBFF;border:1px solid #DDE3EC;
                        border-radius:8px;padding:10px;margin-left:10px;">
              {"".join(f'''<div style="background:white;border:1px solid #DDE3EC;border-radius:6px;
                               padding:6px 8px;margin-bottom:6px;display:flex;
                               justify-content:space-between;align-items:center;">
                <span style="font-size:9px;color:#7F8C8D;font-weight:bold;">{l}</span>
                <span style="font-size:15px;font-weight:900;color:#1A3A5C;">{v}</span>
              </div>''' for l,v in [
                  ("عدد الطلاب",str(n)),("أعلى درجة",f"{high:.0f}"),("أقل درجة",f"{low:.0f}"),
                  ("متوسط الدرجات",f"{avg*mx/100:.1f}"),("نسبة التحصيل",f"{avg:.1f}%"),("مجموع الدرجات",f"{total:.0f}")])}
            </div>
            <div style="flex:1;">
              <table style="width:100%;border-collapse:collapse;font-family:Tahoma,Arial,sans-serif;">
                <thead><tr style="background:#1A3A5C;">
                  <th style="padding:8px;color:white;font-size:10px;">المستوى</th>
                  <th style="padding:8px;color:white;font-size:10px;">النطاق</th>
                  <th style="padding:8px;color:white;font-size:10px;">عدد الطلاب</th>
                  <th style="padding:8px;color:white;font-size:10px;">النسبة</th>
                </tr></thead>
                <tbody>{grade_rows}</tbody>
              </table>
            </div>
          </div>
          <div style="margin:14px 20px 0;">
            <div style="background:#EBF5FB;border-radius:5px;padding:5px 12px;text-align:center;
                        font-size:10px;color:#2471A3;font-weight:bold;margin-bottom:8px;">
              رسم بياني — توزيع الطلاب على التقديرات
            </div>
            <div style="display:flex;gap:20px;align-items:flex-start;">
              <div style="flex:1;">{bars_vis}</div>
              <div style="flex:1;display:flex;flex-wrap:wrap;gap:10px;justify-content:center;">{donuts}</div>
            </div>
          </div>
          <div style="position:absolute;bottom:0;left:0;right:0;padding:10px 28px;
                      border-top:1px solid #DDE3EC;display:flex;
                      justify-content:space-between;align-items:center;">
            <div style="text-align:center;width:180px;">
              <div style="font-size:9px;color:#7F8C8D;margin-bottom:4px;">معلم المادة</div>
              <div style="font-size:10px;font-weight:bold;color:#1A3A5C;margin-bottom:12px;">{tchr or '_______________'}</div>
              <div style="border-top:1px solid #BDC3C7;margin:0 16px;"></div>
            </div>
            <div style="font-size:8px;color:#BDC3C7;">www.edu-forms.com</div>
            <div style="text-align:center;width:180px;">
              <div style="font-size:9px;color:#7F8C8D;margin-bottom:4px;">مدير المدرسة</div>
              <div style="font-size:10px;font-weight:bold;color:#1A3A5C;margin-bottom:12px;">{prin or '_______________'}</div>
              <div style="border-top:1px solid #BDC3C7;margin:0 16px;"></div>
            </div>
          </div>
        </div>"""

    return f"""<!DOCTYPE html><html lang="ar" dir="rtl"><head><meta charset="UTF-8">\n<title>تحليل نتائج الطلاب</title>\n<style>\n@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap');\n*{{box-sizing:border-box;margin:0;padding:0;}}\nbody{{font-family:'Cairo',Tahoma,Arial,sans-serif;background:#f0f0f0;direction:rtl;color:#1A2332;}}\n.page{{width:794px;min-height:1123px;margin:0 auto 30px;background:white;\nposition:relative;box-shadow:0 4px 30px rgba(0,0,0,0.15);\npage-break-after:always;overflow:hidden;}}\n.no-print{{background:#1A3A5C;padding:12px 24px;display:flex;align-items:center;\ngap:12px;position:sticky;top:0;z-index:99;}}\n@media print{{body{{background:white;}}.page{{box-shadow:none;margin:0;width:100%;min-height:100vh;}}\n.no-print{{display:none;}}}}\n</style></head><body>\n<div class="no-print">
  <button onclick="window.print()" style="background:#27AE60;color:white;border:none;
     padding:8px 18px;border-radius:6px;cursor:pointer;font-size:13px;font-weight:bold;
     font-family:Cairo,Tahoma,Arial;">🖨️ طباعة / حفظ PDF</button>
  <span style="color:rgba(255,255,255,0.7);font-size:12px;">
    في نافذة الطباعة فعّل خيار «خلفيات الرسومات» للحصول على الألوان الكاملة
  </span>
</div>
{pages}</body></html>"""


# ════════════════════════════════════════════════════════════
# ── تصدير Word (.docx) ────────────────────────────────────
# ════════════════════════════════════════════════════════════
def _ga_export_word(students, output_path, sel_subject="الكل"):
    """يصدّر تقرير تحليل النتائج بصيغة Word .docx"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        raise ImportError("مكتبة python-docx غير مثبّتة\nقم بتثبيتها: pip install python-docx")

    if not students:
        raise ValueError("لا توجد بيانات للتصدير")

    first = students[0]
    hdr   = _GA_HEADER_DATA
    subjects_list = (
        sorted(set(sub["subject"] for s in students for sub in s.get("subjects", [])))
        if sel_subject == "الكل" else [sel_subject]
    )

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.insert(0, bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)

    def cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color.lstrip('#'))
        tcPr.append(shd)

    def add_cell_text(cell, text, bold=False, size=10, color=None, align='center'):
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(para)
        run = para.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        if color:
            run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))

    doc = DocxDocument()
    # إعداد الصفحة A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    cl  = first.get('class', '') or hdr.get('grade_class', '')
    tm  = first.get('term', '')  or hdr.get('term', '')
    sy  = first.get('school_year', '') or hdr.get('school_year', '')
    sch = first.get('school', '') or hdr.get('school', '')
    pr  = first.get('period', '') or hdr.get('period', '')
    tchr = hdr.get('teacher', '')
    prin = hdr.get('principal', '')
    reg  = hdr.get('region', '')

    for subj_idx, subj_name in enumerate(subjects_list):
        pairs = []
        for s in students:
            for sub in s.get("subjects", []):
                if sub["subject"] == subj_name:
                    pairs.append((sub["score"], sub["max_score"]))
                    break
        if not pairs:
            continue

        n  = len(pairs)
        mx = pairs[0][1]
        pcts  = [sc / m * 100 if m > 0 else 0 for sc, m in pairs]
        avg   = sum(pcts) / len(pcts)
        high  = max(sc for sc, m in pairs)
        low   = min(sc for sc, m in pairs)
        total = sum(sc for sc, m in pairs)
        gc = {g[0]: 0 for g in _GA_GRADES}
        for p in pcts:
            lbl, *_ = _ga_grade(p)
            gc[lbl] += 1

        if subj_idx > 0:
            doc.add_page_break()

        # ── العنوان الرئيسي ──
        title_p = doc.add_paragraph()
        set_rtl(title_p)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(f"تحليل نتائج مادة [{subj_name}]")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = 'Arial'
        title_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        # ── معلومات المدرسة ──
        info_tbl = doc.add_table(rows=1, cols=3)
        info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_tbl.style = 'Table Grid'
        w_col = [Cm(5.5), Cm(5.5), Cm(5.5)]
        for i, (lbl, val) in enumerate([
            ("المرحلة / الصف",    cl),
            ("السنة / الفصل الدراسي", f"{tm} / {sy}هـ"),
            ("درجة القياس",       str(int(mx))),
        ]):
            cell = info_tbl.rows[0].cells[i]
            cell.width = w_col[i]
            cell_bg(cell, "EBF5FB")
            add_cell_text(cell, f"{lbl}\n{val}", bold=True, size=10, color="1A3A5C")
        doc.add_paragraph()

        # ── جدول الإحصائيات التفصيلية ──
        stats_title = doc.add_paragraph()
        set_rtl(stats_title)
        stats_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = stats_title.add_run("الإحصائيات التفصيلية")
        sr.bold = True; sr.font.size = Pt(12); sr.font.name = 'Arial'
        sr.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

        # جدول الملخص السريع (6 بطاقات)
        summ_tbl = doc.add_table(rows=2, cols=3)
        summ_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        summ_tbl.style = 'Table Grid'
        summary_items = [
            ("عدد الطلاب", str(n)),
            ("أعلى درجة",  f"{high:.0f}"),
            ("أقل درجة",   f"{low:.0f}"),
            ("متوسط الدرجات", f"{avg*mx/100:.1f}"),
            ("نسبة التحصيل",  f"{avg:.1f}%"),
            ("مجموع الدرجات", f"{total:.0f}"),
        ]
        for idx, (lbl, val) in enumerate(summary_items):
            row_i, col_i = divmod(idx, 3)
            cell = summ_tbl.rows[row_i].cells[col_i]
            cell_bg(cell, "F8FBFF")
            add_cell_text(cell, f"{lbl}\n{val}", bold=True, size=11, color="1A3A5C")
        doc.add_paragraph()

        # جدول التقديرات
        grade_tbl = doc.add_table(rows=len(_GA_GRADES) + 1, cols=4)
        grade_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        grade_tbl.style = 'Table Grid'
        for j, hdr_txt in enumerate(["المستوى", "النطاق", "عدد الطلاب", "النسبة %"]):
            cell = grade_tbl.rows[0].cells[j]
            cell_bg(cell, "1A3A5C")
            add_cell_text(cell, hdr_txt, bold=True, size=10, color="FFFFFF")
        for i, (lbl, lo, hi, col, bg) in enumerate(_GA_GRADES):
            cnt  = gc[lbl]
            pct_s = f"{cnt/n*100:.0f}%"
            rng_s = f"{int(lo)}–{int(hi) if hi<100 else int(mx)}"
            row_data = [lbl, rng_s, str(cnt), pct_s]
            row_bg_hex = bg.lstrip('#')
            for j, val in enumerate(row_data):
                cell = grade_tbl.rows[i + 1].cells[j]
                cell_bg(cell, row_bg_hex if j == 0 else ("F4F6F7" if i % 2 == 0 else "FFFFFF"))
                is_grade_col = (j == 0)
                add_cell_text(cell, val, bold=is_grade_col, size=10,
                              color=col.lstrip('#') if is_grade_col else "2C3E50")
        doc.add_paragraph()

        # التوقيعات
        sig_tbl = doc.add_table(rows=1, cols=3)
        sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_tbl.style = 'Table Grid'
        for j, (role, name) in enumerate([
            ("معلم المادة", tchr or "_______________"),
            ("", ""),
            ("مدير المدرسة", prin or "_______________"),
        ]):
            cell = sig_tbl.rows[0].cells[j]
            cell_bg(cell, "FFFFFF")
            if role:
                add_cell_text(cell, f"{role}\n{name}", size=10, color="1A3A5C")

        # ملاحظة
        note_p = doc.add_paragraph()
        set_rtl(note_p)
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_p.add_run("www.edu-forms.com — موقع نماذج تعليمية")
        note_run.font.size = Pt(8)
        note_run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

    doc.save(output_path)


def _ga_placeholder_html(msg=None):
    msg = msg or (
        "📊 ارفع ملف نتائج للبدء في التحليل<br><br>"
        "<span style='font-size:12px;color:#95A5A6;line-height:2;'>"
        "📋 Excel .xlsx — ملفات إشعار فترة نور (متعدد الأوراق)<br>"
        "📄 PDF — تقارير نور (يستخرج الدرجات تلقائياً حتى مع الخطوط المشفرة)<br>"
        "📊 CSV — اسم الطالب / المادة / الدرجة / النهاية العظمى"
        "</span>"
    )
    return f"""<!DOCTYPE html><html lang="ar" dir="rtl"><head><meta charset="UTF-8">\n<style>body{{font-family:Tahoma,Arial,sans-serif;background:#F0F4F8;display:flex;\nalign-items:center;justify-content:center;height:100vh;margin:0;direction:rtl;}}\n.ph{{text-align:center;color:#7F8C8D;}}.ph .icon{{font-size:64px;opacity:0.4;margin-bottom:16px;}}\n.ph h3{{font-size:17px;color:#1A3A5C;opacity:0.7;margin-bottom:10px;}}\n</style></head><body><div class="ph">
<div class="icon">📊</div><h3>تحليل نتائج الطلاب</h3>
<p style="font-size:14px;">{msg}</p></div></body></html>"""


# ════════════════════════════════════════════════════════════
# ── نافذة تعديل بيانات الترويسة ──────────────────────────
# ════════════════════════════════════════════════════════════
def _ga_open_header_editor(parent_root, on_save=None):
    """نافذة منبثقة لتعديل بيانات الترويسة (المدرسة / الفصل / المعلم / المدير)"""
    win = tk.Toplevel(parent_root)
    win.title("✏️ تعديل بيانات الترويسة")
    win.geometry("480x420")
    win.resizable(False, False)
    win.transient(parent_root)
    win.grab_set()

    # عنوان
    hdr_frame = tk.Frame(win, bg="#1A3A5C", height=46)
    hdr_frame.pack(fill="x"); hdr_frame.pack_propagate(False)
    tk.Label(hdr_frame, text="✏️ بيانات الترويسة والتقرير",
             bg="#1A3A5C", fg="white",
             font=("Tahoma", 11, "bold")).pack(side="right", padx=14, pady=12)

    body = tk.Frame(win, bg="white", padx=20, pady=16)
    body.pack(fill="both", expand=True)

    fields = [
        ("school",      "اسم المدرسة"),
        ("grade_class", "المرحلة / الصف"),
        ("school_year", "العام الدراسي"),
        ("term",        "الفصل الدراسي"),
        ("period",      "نوع الاختبار / الإشعار"),
        ("region",      "الإدارة التعليمية"),
        ("teacher",     "اسم معلم المادة"),
        ("principal",   "اسم مدير المدرسة"),
    ]

    vars_map = {}
    for i, (key, label) in enumerate(fields):
        row = tk.Frame(body, bg="white")
        row.pack(fill="x", pady=4)
        tk.Label(row, text=f"{label}:", width=20, anchor="e",
                 bg="white", fg="#1A3A5C",
                 font=("Tahoma", 9, "bold")).pack(side="right", padx=(0, 6))
        var = tk.StringVar(value=_GA_HEADER_DATA.get(key, ""))
        ttk.Entry(row, textvariable=var, width=28,
                  font=("Tahoma", 9)).pack(side="right", fill="x", expand=True)
        vars_map[key] = var

    def _save():
        for key, var in vars_map.items():
            _GA_HEADER_DATA[key] = var.get().strip()
        if on_save:
            on_save()
        win.destroy()

    btn_row = tk.Frame(win, bg="#F5F7FA", pady=10)
    btn_row.pack(fill="x")
    ttk.Button(btn_row, text="💾 حفظ البيانات",
               command=_save).pack(side="right", padx=14)
    ttk.Button(btn_row, text="إلغاء",
               command=win.destroy).pack(side="right", padx=6)


# ════════════════════════════════════════════════════════════
# ── دالة بناء التبويب (داخل AppGUI) ─────────────────────
# ════════════════════════════════════════════════════════════
